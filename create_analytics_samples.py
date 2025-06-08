#!/usr/bin/env python3
"""
Create enhanced sample documents optimized for analytics testing
Includes rich temporal, geographic, and sentiment content
"""

import os
from pathlib import Path
from datetime import datetime, timedelta


def create_analytics_optimized_samples():
    """Create sample documents rich in temporal, geographic, and sentiment data"""

    # Create sample directory
    sample_dir = Path("data/analytics_samples")
    sample_dir.mkdir(parents=True, exist_ok=True)

    # Enhanced intelligence reports with rich analytics content
    analytics_samples = {
        "timeline_operation": f"""INTELLIGENCE REPORT - OPERATION CHRONOS
Date: {datetime.now().strftime('%B %d, %Y')}
Classification: SECRET
Analyst: Senior Intelligence Officer Sarah Chen

Subject: Timeline Analysis - EASTERN SYNDICATE Operations

EXECUTIVE SUMMARY:
Multi-source intelligence indicates escalating activities by the EASTERN SYNDICATE 
across multiple geographic regions. Timeline analysis reveals critical acceleration 
in operations beginning January 15, 2025.

TEMPORAL INTELLIGENCE:
- January 15, 2025: Initial contact between VLADIMIR PETROV and DMITRI VOLKOV in MOSCOW
- February 3, 2025: Financial transfers totaling $2.5 million through SWISS NATIONAL BANK
- February 20, 2025: Suspicious meeting in LONDON between PETROV and unknown associates
- March 5, 2025: Intelligence intercept indicates "Operation Thunder" planning phase
- March 18, 2025: CRIMSON ENTERPRISES registered in CAYMAN ISLANDS
- April 2, 2025: Large-scale cryptocurrency transactions detected (450 BTC)
- April 15, 2025: Surveillance reports increased activity in AMSTERDAM and DUBAI
- May 1, 2025: Critical intelligence suggests imminent operation launch
- May 20, 2025: Emergency funding of $15.7 million identified
- June 5, 2025: Current threat assessment upgraded to CRITICAL

THREAT ASSESSMENT:
This represents an immediate and critical threat to national security. 
The accelerating timeline suggests urgent action is required within 48 hours.
Recommend immediate surveillance and potential disruption operations.

NEXT REVIEW: Daily briefings required until threat neutralized
PRIORITY: IMMEDIATE ACTION REQUIRED""",

        "geographic_intelligence": f"""MULTI-SOURCE INTELLIGENCE BRIEFING
Date: {datetime.now().strftime('%B %d, %Y')}
Classification: TOP SECRET//NOFORN
Source: GEOGRAPHIC INTELLIGENCE FUSION CENTER

Subject: Global Network Analysis - SHADOW FINANCIAL NETWORK

GEOGRAPHIC FOOTPRINT ANALYSIS:

EUROPEAN OPERATIONS HUB:
- LONDON, UNITED KINGDOM: Primary financial coordination center
- AMSTERDAM, NETHERLANDS: Logistics and transportation hub  
- ZURICH, SWITZERLAND: Money laundering operations through ALPINE BANK
- MONACO: High-value asset storage and luxury real estate investments
- BERLIN, GERMANY: Technical operations and cyber infrastructure

EASTERN EUROPE CORRIDOR:
- MOSCOW, RUSSIA: Strategic planning and leadership coordination
- BUCHAREST, ROMANIA: Data center operations and digital infrastructure
- SOFIA, BULGARIA: Transportation and smuggling route management
- KIEV, UKRAINE: Regional coordination and local asset recruitment

MIDDLE EAST & ASIA OPERATIONS:
- DUBAI, UAE: Regional financial hub and gold trading operations
- SINGAPORE: Cryptocurrency exchange and digital asset management
- HONG KONG: Asian market penetration and legitimate business fronts
- ISTANBUL, TURKEY: Critical transportation hub linking Europe and Asia

OFFSHORE TERRITORIES:
- CAYMAN ISLANDS: Shell company registrations and tax avoidance
- BRITISH VIRGIN ISLANDS: Complex ownership structures
- PANAMA CITY, PANAMA: Document creation and corporate services
- SEYCHELLES: Offshore banking and asset protection

WESTERN OPERATIONS:
- NEW YORK, USA: Financial market manipulation and legitimate business integration
- MIAMI, FLORIDA: Real estate investments and money laundering
- VANCOUVER, CANADA: Real estate market exploitation
- TORONTO, CANADA: Financial services infiltration

THREAT LEVEL BY REGION:
- CRITICAL: LONDON, MOSCOW, DUBAI (Primary command centers)
- HIGH: AMSTERDAM, ZURICH, SINGAPORE (Major operational hubs)  
- MEDIUM: All other identified locations (Support operations)

GEOGRAPHIC CORRELATION ANALYSIS:
Pattern analysis indicates coordinated operations across time zones,
suggesting sophisticated command and control structure with regional autonomy.

RECOMMENDATION: Enhanced surveillance on all identified locations with
priority focus on CRITICAL threat regions.""",

        "sentiment_threat_assessment": f"""THREAT ASSESSMENT REPORT
Date: {datetime.now().strftime('%B %d, %Y')}
Classification: SECRET
Threat Level: HIGH
Urgency: IMMEDIATE

Subject: Critical Security Situation - IMMINENT THREAT ANALYSIS

THREAT SUMMARY:
URGENT intelligence indicates CRITICAL and IMMEDIATE threats to infrastructure.
This is an EMERGENCY situation requiring URGENT response within 24 hours.

SITUATION ANALYSIS:
Advanced Persistent Threat group APT-PHANTOM has demonstrated ALARMING escalation
in aggressive cyber operations. Intelligence suggests IMMINENT large-scale attack
on CRITICAL INFRASTRUCTURE including POWER GRIDS and FINANCIAL INSTITUTIONS.

IMMEDIATE CONCERNS:
- URGENT: Zero-day exploits targeting WINDOWS SERVER infrastructure
- CRITICAL: Malware deployment in progress across TELECOMMUNICATIONS networks
- EMERGENCY: Suspected coordination with state-sponsored actors
- ALARMING: Evidence of insider threats within DEFENSE CONTRACTORS

EMOTIONAL INTELLIGENCE ASSESSMENT:
Communications intercepts reveal heightened urgency and aggressive posturing
from threat actors. Language analysis indicates confidence in planned operations
and dismissive attitude toward defensive capabilities.

Recent intercepts include phrases such as:
- "Immediate action required" - suggesting rushed timeline
- "Critical window of opportunity" - indicating time-sensitive operations  
- "Devastating impact expected" - showing malicious intent
- "No turning back" - demonstrating commitment to harmful actions

PSYCHOLOGICAL PROFILE:
Threat actors demonstrate dangerous combination of technical sophistication
and reckless disregard for consequences. This represents an extremely 
concerning threat profile requiring IMMEDIATE intervention.

URGENCY INDICATORS:
Multiple sources confirm IMMEDIATE action required. This is NOT a drill.
Threat assessment indicates CRITICAL situation with potential for SEVERE
consequences if not addressed within next 24-48 hours.

RECOMMENDED ACTIONS:
1. IMMEDIATE elevation to National Security Council
2. URGENT deployment of cyber defense teams
3. CRITICAL infrastructure protection protocols activated
4. EMERGENCY coordination with international partners

This situation requires URGENT attention at the highest levels.
Lives and national security are at IMMEDIATE risk.

Status: ACTIVE MONITORING - URGENT RESPONSE REQUIRED""",

        "multi_dimensional_analysis": f"""COMPREHENSIVE INTELLIGENCE ANALYSIS
Date: {datetime.now().strftime('%B %d, %Y')}
Classification: SECRET//REL TO USA, CAN, GBR, AUS
Analysis Type: MULTI-DIMENSIONAL FUSION

Subject: Comprehensive Assessment - GLOBAL SYNDICATE OPERATIONS

EXECUTIVE SUMMARY:
Comprehensive analysis combining temporal, geographic, and sentiment intelligence
reveals sophisticated international criminal organization with concerning capabilities.

TIMELINE RECONSTRUCTION:
- December 1, 2024: Initial formation of criminal network in EASTERN EUROPE
- January 15, 2025: First international expansion to WESTERN EUROPE
- February 28, 2025: Establishment of NORTH AMERICAN operations
- March 15, 2025: MIDDLE EAST penetration through DUBAI operations
- April 10, 2025: ASIA-PACIFIC expansion via SINGAPORE and HONG KONG
- May 1, 2025: Full global network operational
- Present: Preparing for major coordinated operation

GEOGRAPHIC ANALYSIS:
Network spans five continents with command structure based in:
- Primary: LONDON (Financial), MOSCOW (Strategic), DUBAI (Regional)
- Secondary: SINGAPORE, NEW YORK, AMSTERDAM (Operational hubs)
- Tertiary: Multiple cities across 25+ countries (Support operations)

Key geographic vulnerabilities identified in SWITZERLAND, CAYMAN ISLANDS,
and PANAMA where regulatory oversight appears insufficient.

SENTIMENT & THREAT ANALYSIS:
Communications analysis reveals escalating confidence and aggressive intent.
Recent intercepts show URGENT planning for "final phase" operations.
Psychological assessment indicates dangerous combination of sophistication
and willingness to cause SEVERE harm to achieve objectives.

FINANCIAL INTELLIGENCE:
Total network assets estimated at $2.8 billion across multiple jurisdictions.
Money laundering operations process approximately $150 million monthly.
CRYPTOCURRENCY usage increasing, with 15,000+ BTC under network control.

THREAT ASSESSMENT:
CRITICAL threat to multiple national interests. Network demonstrates:
- IMMEDIATE operational capability
- URGENT timeline for major operations  
- SEVERE potential for economic disruption
- ALARMING level of international coordination

MULTI-SOURCE CORRELATION:
All intelligence sources confirm IMMEDIATE and CRITICAL threat requiring
URGENT response. Timeline analysis suggests operation planned for next 30 days.
Geographic spread makes unilateral response insufficient.

RECOMMENDATIONS:
1. IMMEDIATE international coordination required
2. URGENT asset freezing across all identified jurisdictions
3. CRITICAL infrastructure protection measures
4. Enhanced surveillance on all primary and secondary hubs

This represents one of the most SERIOUS and IMMEDIATE threats identified
in recent intelligence assessments. URGENT action required.""",

        "operational_update": f"""OPERATIONAL INTELLIGENCE UPDATE
Date: {datetime.now().strftime('%B %d, %Y')}
Time: 14:30 UTC
Classification: CONFIDENTIAL
Priority: HIGH

Subject: Operation NIGHTWATCH - Status Update

OPERATIONAL TIMELINE:
- 06:00 UTC: Surveillance team Alpha deployed to FRANKFURT
- 08:15 UTC: Target MARCUS WEBB observed entering DEUTSCHE BANK building
- 09:30 UTC: Suspicious meeting with unknown associate (SUBJECT-47)
- 11:45 UTC: Electronic surveillance indicates financial transaction discussion
- 12:20 UTC: WEBB departed via vehicle (License: DE-AB-5678) toward MUNICH
- 13:15 UTC: Communications intercept reveals reference to "ZURICH operation"
- 14:00 UTC: Enhanced surveillance requested for SWISS NATIONAL BANK

GEOGRAPHIC TRACKING:
Route analysis indicates systematic visits to financial institutions across:
- FRANKFURT, GERMANY (DEUTSCHE BANK headquarters)
- MUNICH, GERMANY (Regional financial center)  
- Projected destination: ZURICH, SWITZERLAND

SENTIMENT ANALYSIS OF COMMUNICATIONS:
Intercepted communications show concerning language patterns:
- Repeated use of "urgent" and "immediate" (indicating time pressure)
- References to "final arrangements" (suggesting completion phase)
- Mentions of "significant consequences" (indicating major operation)

THREAT LEVEL ASSESSMENT:
Based on behavioral analysis and communication patterns, threat level
upgraded from MEDIUM to HIGH. Subject demonstrates concerning urgency
and appears to be coordinating time-sensitive financial operations.

IMMEDIATE ACTIONS REQUIRED:
1. Coordinate with GERMAN FEDERAL POLICE for continued surveillance
2. Alert SWISS AUTHORITIES regarding potential ZURICH operations
3. Request INTERPOL notification for cross-border tracking
4. Enhance monitoring of all associated financial institutions

Next update scheduled: 18:00 UTC or upon significant developments.

Status: ACTIVE OPERATION - ENHANCED SURVEILLANCE APPROVED"""
    }

    print("📊 Creating analytics-optimized intelligence samples...")

    # Create text files
    for filename, content in analytics_samples.items():
        txt_file = sample_dir / f"{filename}.txt"
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"✅ Created: {txt_file}")

    # Create Word documents if available
    try:
        from docx import Document

        print("\n📄 Creating Word document versions...")
        for filename, content in analytics_samples.items():
            doc = Document()

            # Add title with classification
            title = f'{filename.replace("_", " ").title()} Report'
            doc.add_heading(title, 0)

            # Add classification header
            classification_para = doc.add_paragraph()
            classification_para.add_run("CLASSIFICATION: SECRET").bold = True

            # Split content into paragraphs and add to document
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Check if it's a header (all caps or starts with specific patterns)
                    if (para.isupper() and len(para) < 100) or para.startswith(
                            ('EXECUTIVE', 'TIMELINE', 'GEOGRAPHIC', 'THREAT')):
                        doc.add_heading(para.strip(), level=1)
                    else:
                        doc.add_paragraph(para.strip())

            docx_file = sample_dir / f"{filename}.docx"
            doc.save(str(docx_file))
            print(f"✅ Created: {docx_file}")

    except ImportError:
        print("⚠️ python-docx not available, skipping Word document creation")

    # Create summary file with testing instructions
    instructions_file = sample_dir / "ANALYTICS_TESTING_GUIDE.md"

    instructions_content = f"""# Advanced Analytics Testing Guide

## 📊 Sample Documents Created

The following intelligence documents have been optimized for testing advanced analytics:

### 1. **timeline_operation** - Timeline Analysis Testing
- **Focus**: Temporal pattern analysis
- **Content**: Detailed timeline from January to June 2025
- **Tests**: Date extraction, temporal clustering, timeline visualization

### 2. **geographic_intelligence** - Geographic Mapping Testing  
- **Focus**: Geographic pattern analysis
- **Content**: Global network across 20+ cities and 15+ countries
- **Tests**: Location extraction, geographic clustering, hotspot identification

### 3. **sentiment_threat_assessment** - Sentiment Analysis Testing
- **Focus**: Threat level and urgency assessment
- **Content**: High-urgency language with threat indicators
- **Tests**: Sentiment scoring, threat classification, urgency detection

### 4. **multi_dimensional_analysis** - Combined Analytics Testing
- **Focus**: All analytics features simultaneously
- **Content**: Rich temporal, geographic, and sentiment data
- **Tests**: Cross-analytics correlation, comprehensive insights

### 5. **operational_update** - Real-time Intelligence Testing
- **Focus**: Operational timeline and geographic tracking
- **Content**: Hour-by-hour operational intelligence
- **Tests**: Real-time analysis, geographic tracking, sentiment monitoring

## 🧪 Testing Workflow

### Step 1: Upload and Process Documents
1. Go to **Document Upload** tab
2. Upload all sample files (.txt or .docx versions)
3. Set document type to "security_intelligence" or "threat_assessment"
4. Process all documents

### Step 2: Timeline Analysis Testing
1. Navigate to **📅 Timeline Analysis** tab
2. Verify temporal pattern extraction
3. Check timeline visualization
4. Examine monthly activity charts
5. Review key temporal entities

### Step 3: Geographic Mapping Testing
1. Navigate to **🗺️ Geographic Mapping** tab  
2. Verify location extraction and geocoding
3. Check interactive map with markers
4. Review geographic hotspots
5. Examine country-level analysis

### Step 4: Sentiment Analysis Testing
1. Navigate to **📊 Sentiment Analysis** tab
2. Verify threat level classification
3. Check urgency scoring
4. Review sentiment distribution
5. Examine high-priority document flagging

### Step 5: Cross-Analytics Verification
1. Check **Insights Dashboard** for comprehensive view
2. Verify **Relationship Network** includes geographic entities
3. Examine **Entity Analysis** for temporal and geographic entities
4. Test export functionality for all analytics data

## 🎯 Expected Results

### Timeline Analysis
- **15-20 dates** extracted across all documents
- **Monthly activity chart** showing peak activity periods
- **Temporal clusters** around major operation phases
- **Key temporal entities** ranked by importance

### Geographic Mapping
- **25+ locations** identified and geocoded
- **Interactive map** with clustered markers
- **5-10 geographic hotspots** identified
- **Country analysis** showing global distribution

### Sentiment Analysis  
- **Threat levels**: Mix of HIGH, MEDIUM, LOW classifications
- **2-3 documents** flagged as high priority
- **Urgency scores**: Range from 0.1 to 1.0
- **Sentiment distribution**: Primarily negative (intelligence context)

## 🔍 Advanced Testing Scenarios

### Scenario 1: Crisis Response Simulation
1. Process all documents in sequence
2. Use Timeline Analysis to understand escalation
3. Use Geographic Mapping to identify key locations
4. Use Sentiment Analysis to prioritize urgent responses

### Scenario 2: Intelligence Fusion
1. Process documents of different types (mix security/economic/threat)
2. Cross-reference entities across all analytics tabs
3. Build comprehensive intelligence picture
4. Export data for further analysis

### Scenario 3: Real-time Intelligence
1. Process "operational_update" document
2. Track real-time geographic movement
3. Analyze sentiment changes over time
4. Assess threat escalation patterns

## 💡 Pro Tips

- **Geographic Mapping**: Wait for geocoding to complete (may take 30-60 seconds)
- **Timeline Analysis**: Look for patterns in monthly activity charts  
- **Sentiment Analysis**: Pay attention to high-priority document flags
- **Cross-Analytics**: Use insights from one tab to inform analysis in others
- **Export Data**: Use CSV exports for deeper analysis in external tools

## 🎉 Success Indicators

✅ **Timeline**: Clear temporal patterns with accurate date extraction  
✅ **Geographic**: Interactive map with properly geocoded locations  
✅ **Sentiment**: Accurate threat classification and urgency scoring  
✅ **Integration**: Seamless flow between all analytics features  
✅ **Performance**: Smooth processing of all sample documents  

---

**Ready to demonstrate world-class intelligence analytics!** 🚀
"""

    with open(instructions_file, 'w', encoding='utf-8') as f:
        f.write(instructions_content)

    print(f"\n🎉 Analytics samples created in: {sample_dir}")
    print("📋 Files created:")

    for file_path in sample_dir.iterdir():
        if file_path.is_file():
            file_size = file_path.stat().st_size
            print(f"   {file_path.name} ({file_size:,} bytes)")

    print(f"\n📖 Testing guide: {instructions_file}")
    print("\n🚀 Next steps:")
    print("1. Upload these samples in your dashboard")
    print("2. Process all documents")
    print("3. Explore the three new analytics tabs:")
    print("   📅 Timeline Analysis")
    print("   🗺️ Geographic Mapping")
    print("   📊 Sentiment Analysis")


if __name__ == "__main__":
    create_analytics_optimized_samples()