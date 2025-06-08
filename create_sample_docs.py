#!/usr/bin/env python3
"""
Create sample documents in different formats for testing
Run this script to generate test files in various formats
"""

import os
from pathlib import Path


def create_sample_documents():
    """Create sample intelligence documents in different formats"""

    # Create sample_docs directory
    sample_dir = Path("data/sample_docs")
    sample_dir.mkdir(parents=True, exist_ok=True)

    # Sample intelligence report content
    sample_content = {
        "security_report": """INTELLIGENCE REPORT - OPERATION MIDNIGHT
Date: 2025-06-08
Classification: SECRET
Source: FIELD INTELLIGENCE UNIT ALPHA

Subject: Cross-Border Financial Investigation - SHADOW NETWORK

Executive Summary:
Our multi-source intelligence analysis has identified a sophisticated financial network 
operating across three continents. Key findings indicate coordination between 
EASTERN FINANCIAL GROUP (registered in CYPRUS) and PACIFIC TRADING LLC (SINGAPORE).

Key Entities Identified:
- Primary Coordinator: DMITRI VOLKOV (Russian national, DOB: 1978-03-15)
- Financial Controller: SARAH CHEN (Canadian passport, currently in HONG KONG)
- Operations Manager: CARLOS RODRIGUEZ (Mexican national, operates from PANAMA CITY)

Financial Analysis:
Total transaction volume: $847 million over 18-month period
Primary accounts: OFFSHORE BANK OF SEYCHELLES, CARIBBEAN TRUST FUND
Cryptocurrency involvement: 450 BTC traced through mixing services
Real estate investments: $125 million in LONDON, DUBAI, and VANCOUVER

Geographic Pattern:
- Source countries: EASTERN EUROPE (Romania, Bulgaria, Ukraine)
- Transit hubs: ISTANBUL, AMSTERDAM, PANAMA CITY, SINGAPORE
- Destination markets: NORTH AMERICA, WESTERN EUROPE, AUSTRALIA

Recommended Actions:
1. Coordinate with INTERPOL for simultaneous operations
2. Request asset freezing through FINANCIAL CRIMES COMMISSION
3. Enhanced surveillance on key transit points
4. Cryptocurrency tracing through BLOCKCHAIN ANALYSIS UNIT

Next Review: June 15, 2025
Contact: AGENT JENNIFER MARTINEZ, FINANCIAL INTELLIGENCE DIVISION""",

        "threat_assessment": """CYBER THREAT ASSESSMENT REPORT
Date: 2025-06-08
Classification: TOP SECRET//NOFORN
Prepared by: CYBER INTELLIGENCE CENTER

Subject: Advanced Persistent Threat - OPERATION DIGITAL PHANTOM

Threat Overview:
Nation-state actor APT-GHOST has significantly escalated cyber operations targeting 
CRITICAL INFRASTRUCTURE across NORTH AMERICA and WESTERN EUROPE.

Attribution Analysis:
High confidence attribution to FOREIGN ADVERSARY X based on:
- Technical indicators consistent with known APT-GHOST tools
- Infrastructure overlap with previous campaigns
- Timing correlation with geopolitical tensions

Target Profile:
Primary Targets:
- POWER GRID operators in TEXAS, CALIFORNIA, and ONTARIO
- FINANCIAL INSTITUTIONS: GLOBAL TRUST BANK, INTERNATIONAL COMMERCE BANK
- TELECOMMUNICATIONS: METRO COMMUNICATIONS, ATLANTIC TELECOM

Technical Analysis:
Malware Family: PHANTOM-TOOLKIT v3.2
- Zero-day exploits targeting WINDOWS SERVER 2019/2022
- Custom backdoors with advanced persistence mechanisms
- Data exfiltration capabilities: 50GB+ per compromised system

Attack Vectors:
1. SPEAR PHISHING targeting IT ADMINISTRATORS and EXECUTIVES
2. SUPPLY CHAIN compromise through SOFTWARE VENDORS
3. WATERING HOLE attacks on industry-specific websites

Risk Assessment:
Impact: CRITICAL - Potential for widespread infrastructure disruption
Probability: HIGH - Active preparations observed
Timeframe: IMMEDIATE - Operations may commence within 30 days

Contact: DR. MICHAEL THOMPSON, SENIOR CYBER THREAT ANALYST""",

        "economic_intelligence": """ECONOMIC INTELLIGENCE BRIEFING
Date: 2025-06-08
Classification: CONFIDENTIAL
Analyst: SENIOR ECONOMIST JENNIFER PATEL

Subject: Market Manipulation Investigation - TECHNOLOGY SECTOR

Investigation Summary:
Coordinated market manipulation scheme targeting QUANTUM TECHNOLOGIES stock price 
detected through advanced trading pattern analysis. Estimated market impact: $4.2 billion.

Key Players Identified:
Primary Orchestrator: MARCUS WEBB (former SEC investigator, current consultant)
- Current employer: APEX CAPITAL GROUP (hedge fund, $15B assets under management)
- Known associates: VLADIMIR PETROV, SARAH MITCHELL (quantitative analysts)

Financial Entities:
- HEDGE FUND ALPHA: $2.8 billion in suspicious positions
- DIGITAL VENTURES LLC: Offshore entity, $500 million in coordinated trades
- PACIFIC INVESTMENT GROUP: Singapore-based, unusual trading patterns

Target Company Analysis:
QUANTUM TECHNOLOGIES (NASDAQ: QTECH)
- Market cap: $45 billion (pre-manipulation)
- Artificial price inflation: 34% over 6-month period
- Merger announcement timing: Coordinated with large position builds

Trading Pattern Analysis:
Suspicious Activity Timeline:
- T-30 days: Large call option purchases ($180 million notional)
- T-15 days: Coordinated buying across multiple accounts
- T-5 days: Insider trading indicators (99.7% confidence)
- T-Day: Merger announcement, immediate profit-taking

Technology Analysis:
High-frequency trading algorithms detected:
- Co-location advantages exploited at NYSE, NASDAQ
- Social media sentiment manipulation through bot networks

Legal Analysis:
Potential Charges:
- Securities fraud (15 USC §78j(b))
- Market manipulation (15 USC §78i(a)(2))
- Money laundering (18 USC §1956)

Evidence Collection:
- Trading records: 847,000 transactions analyzed
- Communication intercepts: 12,500 messages (court-authorized)
- Financial records: 89 bank accounts across 7 jurisdictions

Distribution: FINANCIAL INTELLIGENCE LEADERSHIP ONLY"""
    }

    # Create TXT files
    print("📄 Creating text documents...")
    for doc_type, content in sample_content.items():
        txt_file = sample_dir / f"{doc_type}.txt"
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"✅ Created: {txt_file}")

    # Create DOCX files if python-docx is available
    try:
        from docx import Document

        print("\n📄 Creating Word documents...")
        for doc_type, content in sample_content.items():
            doc = Document()

            # Add title
            title = doc.add_heading(f'{doc_type.replace("_", " ").title()} Report', 0)

            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())

            docx_file = sample_dir / f"{doc_type}.docx"
            doc.save(str(docx_file))
            print(f"✅ Created: {docx_file}")

    except ImportError:
        print("⚠️ python-docx not available, skipping Word document creation")
        print("💡 Run: pip install python-docx")

    # Create a simple CSV file for testing
    csv_content = """Entity,Type,Document,Confidence
DMITRI VOLKOV,PERSON,security_report,0.95
EASTERN FINANCIAL GROUP,ORG,security_report,0.89
QUANTUM TECHNOLOGIES,ORG,economic_intelligence,0.94
MARCUS WEBB,PERSON,economic_intelligence,0.91
APT-GHOST,ORG,threat_assessment,0.88
CYBER INTELLIGENCE CENTER,ORG,threat_assessment,0.92"""

    csv_file = sample_dir / "entity_summary.csv"
    with open(csv_file, 'w', encoding='utf-8') as f:
        f.write(csv_content)
    print(f"✅ Created: {csv_file}")

    print(f"\n🎉 Sample documents created in: {sample_dir}")
    print("📁 Available files:")

    for file_path in sample_dir.iterdir():
        if file_path.is_file():
            file_size = file_path.stat().st_size
            print(f"   {file_path.name} ({file_size:,} bytes)")


if __name__ == "__main__":
    create_sample_documents()